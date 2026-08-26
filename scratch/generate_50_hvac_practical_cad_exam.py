import json
import os
import shutil
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------------------------------------------------------------
# 1. DEFINE 50 DETAILED PRACTICAL HVAC ESSAY & FIELD EXECUTION QUESTIONS
# ---------------------------------------------------------------------

hvac_tasks = [
    # Lắp đặt giàn nóng (Outdoor Unit VRV/VRF Installation) - 5 câu
    {
        "cat": "Lắp đặt giàn nóng VRV/VRF",
        "q": "Dựa vào Bản vẽ Shop Drawing CAD hệ thống ĐHTG (Mã VY.BDT-ĐH-1002), hãy nêu quy trình giám sát và nghiệm thu lắp đặt cụm dàn nóng VRV đặt trên bệ bê tông sân mái: Yêu cầu về khoảng cách tản nhiệt tối thiểu, lò xo giảm chấn và độ cao bệ so với mặt sàn mái để phòng chống ngập nước?",
        "options": [
            "Bệ bê tông cao ≥150mm, lót cao su/lò xo giảm chấn, cách tường sau ≥300mm, cách mặt trước ≥1000mm",
            "Đặt trực tiếp dàn nóng xuống sàn mái không cần lò xo giảm chấn",
            "Đặt dàn nóng sát tường sau 50mm để tiết kiệm diện tích",
            "Bệ bê tông cao 50mm, không cần thoát nước đọng"
        ],
        "correct": 0,
        "cad_title": "SHOP DRAWING DETAILED INSTALLATION OF OUTDOOR VRV UNIT ON ROOF SLAB",
        "tag": "HVAC-CAD-01"
    },
    {
        "cat": "Lắp đặt giàn nóng VRV/VRF",
        "q": "Theo tiêu chuẩn thi công ĐHTG công trình Vũ Yên Hải Phòng, khi ghép nối 2 dàn nóng đơn thành một cụm dàn nóng kép VRV 30HP, quy định vặn xiết bu-lông chân máy và khoảng cách giữa 2 dàn nóng đơn là bao nhiêu?",
        "options": [
            "Khoảng cách giữa 2 dàn nóng đơn ≥20mm (hoặc theo chỉ định hãng), dùng bu-lông nở M12/M16 xiết chặt lò xo giảm chấn",
            "Đặt sát dàn nóng dính liền nhau không có khe hở",
            "Dùng dây thép buộc chân dàn nóng vào bệ bê tông",
            "Khoảng cách giữa 2 dàn nóng đơn phải ≥2000mm"
        ],
        "correct": 0,
        "cad_title": "DUAL VRV OUTDOOR UNIT INTERCONNECTION & VIBRATION ISOLATOR SHOP DETAIL",
        "tag": "HVAC-CAD-02"
    },
    {
        "cat": "Lắp đặt giàn nóng VRV/VRF",
        "q": "Trong bản vẽ chi tiết lắp đặt dàn nóng VRV, bẫy dầu (Oil Trap) trên đường ống hơi hút dâng đứng được yêu cầu thi công ở khoảng cách dâng đứng bao nhiêu mét?",
        "options": [
            "Đặt bẫy dầu cứ mỗi 6m - 10m độ cao dâng đứng (khi dàn nóng đặt cao hơn dàn lạnh)",
            "Không cần làm bẫy dầu trên đường dâng đứng",
            "Chỉ làm bẫy dầu ở ngay chân dàn nóng",
            "Mỗi 1m dâng đứng phải làm 1 bẫy dầu"
        ],
        "correct": 0,
        "cad_title": "VERTICAL SUCTION RISER OIL TRAP DETAIL ON VRV PIPING",
        "tag": "HVAC-CAD-03"
    },
    {
        "cat": "Lắp đặt giàn nóng VRV/VRF",
        "q": "Nghiệm thu hướng thổi gió tản nhiệt dàn nóng VRV trên bệ mái: Nếu xung quanh có tường bao lửng chắn gió tản nhiệt, giải pháp kỹ thuật bắt buộc trên bản vẽ Shop là gì?",
        "options": [
            "Lắp chụp hướng dòng gió thổi đứng (Air Discharge Duct Hood) đẩy không khí nóng vọt qua khỏi đỉnh tường bao",
            "Dừng máy nén khi trời nắng nóng",
            "Tháo bỏ toàn bộ vỏ dàn nóng",
            "Lắp quạt cây thổi vào dàn nóng"
        ],
        "correct": 0,
        "cad_title": "DISCHARGE AIR HOOD EXTENSION FOR OUTDOOR CONDENSER UNIT SHOP DETAIL",
        "tag": "HVAC-CAD-04"
    },
    {
        "cat": "Lắp đặt giàn nóng VRV/VRF",
        "q": "Kiểm tra hệ thống tiếp địa và cấp nguồn điện cho cụm dàn nóng VRV 380V/3P/50Hz: Yêu cầu kích thước dây cáp điện và công tắc cách ly (Isolator Switch) ngoài trời là gì?",
        "options": [
            "Cáp điện chống cháy/chống ẩm đi trong ống ruột gà lõi thép bọc nhựa, lắp Isolator chống nước IP65 ngay tại chân dàn nóng",
            "Đi dây điện trần tiếp xúc trực tiếp ngoài trời",
            "Dùng phích cắm dân dụng 2 chấu",
            "Không cần lắp Aptomat/Isolator bảo vệ ngoài trời"
        ],
        "correct": 0,
        "cad_title": "OUTDOOR ISOLATOR SWITCH & ELECTRICAL CONNECTIONS SHOP DETAIL",
        "tag": "HVAC-CAD-05"
    },

    # Lắp đặt giàn lạnh (Indoor Unit FCU Cassette / Ducted) - 5 câu
    {
        "cat": "Lắp đặt giàn lạnh FCU/Cassette",
        "q": "Dựa vào Bản vẽ Shop Drawing VY.BDT-ĐH-1001, khi treo dàn lạnh Cassette 4 hướng thổi lên trần bê tông, quy cách ty treo và đai ốc chống trôi (Double Nut) được quy định như thế nào?",
        "options": [
            "Dùng ty ren M10/M12 mạ kẽm, nở sắt dầm bê tông, sử dụng đai ốc kép và đệm vênh khóa chặt chân treo dàn lạnh",
            "Dùng dây thép dẻo phi 4 treo dàn lạnh",
            "Dùng 1 đai ốc thả lỏng không cần siết chặt",
            "Khoan đinh vít nở nhựa vào bê tông"
        ],
        "correct": 0,
        "cad_title": "4-WAY CASSETTE INDOOR UNIT CEILING HANGER & THREADED ROD SHOP DETAIL",
        "tag": "HVAC-CAD-06"
    },
    {
        "cat": "Lắp đặt giàn lạnh FCU/Cassette",
        "q": "Nghiệm thu lắp đặt dàn lạnh giấu trần nối ống gió (Ducted FCU): Yêu cầu kỹ thuật về khớp nối mềm (Canvas Connection) nối giữa miệng thổi dàn lạnh với box gió là gì?",
        "options": [
            "Khớp nối mềm bằng vải bạt bọc cách nhiệt, chiều dài 150-200mm, triệt tiêu hoàn toàn độ rung từ quạt FCU truyền vào ống gió",
            "Nối trực tiếp tôn phẳng cứng không có khoảng hở",
            "Khớp nối mềm dài 1000mm rủ võng",
            "Dùng băng dính nilon quấn quanh"
        ],
        "correct": 0,
        "cad_title": "DUCTED FCU CANVAS FLEXIBLE JOINT & PLENUM BOX SHOP DETAIL",
        "tag": "HVAC-CAD-07"
    },
    {
        "cat": "Lắp đặt giàn lạnh FCU/Cassette",
        "q": "Khi nghiệm thu đường ống thoát nước xả dàn lạnh FCU, bẫy nước (Con densate Drain Trap) được lắp đặt nhằm mục đích gì và độ sâu bẫy nước tối thiểu là bao nhiêu?",
        "options": [
            "Ngăn mùi hôi từ trục thoát nước ngược vào phòng và ngăn áp suất quạt hút ngược nước xả; độ sâu bẫy nước ≥50mm",
            "Tăng tốc độ chảy của nước ngưng",
            "Giảm nhiệt độ nước xả",
            "Độ sâu bẫy nước chỉ cần 5mm"
        ],
        "correct": 0,
        "cad_title": "CONDENSATE DRAIN PIPE U-TRAP & DRAINAGE SLOPE SHOP DETAIL",
        "tag": "HVAC-CAD-08"
    },
    {
        "cat": "Lắp đặt giàn lạnh FCU/Cassette",
        "q": "Trên bản vẽ Shop trần hoàn thiện, khoảng cách khoảng hở tối thiểu giữa miệng gió hồi (Return Grille) và miệng gió thổi (Supply Diffuser) của dàn lạnh giấu trần để tránh ngắn mạch không khí là bao nhiêu?",
        "options": [
            "Khoảng cách tối thiểu ≥1500mm (1.5m) hoặc bố trí hai phía đối diện phòng",
            "Bố trí sát nhau 100mm",
            "Lắp chung 1 lỗ trần",
            "Không quan tâm khoảng cách"
        ],
        "correct": 0,
        "cad_title": "SUPPLY & RETURN AIR DIFFUSER SPACING ON RECESSED CEILING DETAIL",
        "tag": "HVAC-CAD-09"
    },
    {
        "cat": "Lắp đặt giàn lạnh FCU/Cassette",
        "q": "Quy trình nghiệm thu cửa thăm trần (Access Door) bảo dưỡng dàn lạnh giấu trần FCU: Kích thước tiêu chuẩn cửa thăm và vị trí mở cửa thăm?",
        "options": [
            "Kích thước tối thiểu 450x450mm hoặc 600x600mm, đặt ngay bên dưới vị trí động cơ quạt, hộp điện và van tiết lưu EEV",
            "Kích thước 100x100mm ở góc phòng",
            "Không cần làm cửa thăm trần",
            "Lắp cửa thăm kín không mở được"
        ],
        "correct": 0,
        "cad_title": "CEILING ACCESS DOOR POSITION FOR FCU MAINTENANCE SHOP DETAIL",
        "tag": "HVAC-CAD-10"
    },

    # Hàn ống đồng (Copper Brazing & Nitrogen Purging) - 5 câu
    {
        "cat": "Hàn ống đồng & Thổi khí Nitơ",
        "q": "Tại công trường Hạ Long Xanh, khi kỹ thuật viên thực hiện hàn vảy bạc nối ống đồng phi 28.6mm, lưu lượng và áp suất khí Nitơ khô thổi bảo vệ qua lòng ống là bao nhiêu?",
        "options": [
            "Thổi khí Nitơ liên tục với áp suất nhẹ khoảng 0.02 - 0.05 MPa (0.2 - 0.5 bar) trong suốt quá trình nung hàn",
            "Thổi khí Nitơ áp suất cao 4.0 MPa trong khi hàn",
            "Thổi khí Oxy vào lòng ống khi hàn",
            "Không cần thổi khí Nitơ"
        ],
        "correct": 0,
        "cad_title": "COPPER PIPE BRAZING WITH CONTINUOUS NITROGEN PURGE SHOP DETAIL",
        "tag": "HVAC-CAD-11"
    },
    {
        "cat": "Hàn ống đồng & Thổi khí Nitơ",
        "q": "Nghiệm thu mối hàn ống đồng bằng kính phóng đại hoặc mắt thường: Tiêu chuẩn một mối hàn ngấu đạt chất lượng kỹ thuật là gì?",
        "options": [
            "Vảy hàn điền đầy 100% khe hở mối nối, bề mặt láng mịn tròn đều, lòng ống sạch tuyệt đối không bám muội đen oxy hóa",
            "Mối hàn bám cục xù xì, lòng ống đốm đen vảy cá",
            "Mối hàn chỉ cần dính 30% khe hở",
            "Mối hàn bị cháy thủng ống đồng"
        ],
        "correct": 0,
        "cad_title": "COPPER BRAZED JOINT CROSS-SECTION & ACCEPTANCE CRITERIA DETAIL",
        "tag": "HVAC-CAD-12"
    },
    {
        "cat": "Hàn ống đồng & Thổi khí Nitơ",
        "q": "Khi loe ống đồng (Flaring) để bắt giắc co vào van dàn lạnh, góc loe tiêu chuẩn và mặt nón đầu loe phải đạt yêu cầu gì?",
        "options": [
            "Góc loe 90 độ, nón loe phẳng mịn không gờ nứt, bề mặt tiếp xúc sáng bóng không bị lệch tâm",
            "Nón loe bị gờ mép rách răng cưa",
            "Dùng búa đập bẹp đầu ống",
            "Góc loe 45 độ lệch nón"
        ],
        "correct": 0,
        "cad_title": "COPPER TUBE FLARING & FLARE NUT CONNECTION SHOP DETAIL",
        "tag": "HVAC-CAD-13"
    },
    {
        "cat": "Hàn ống đồng & Thổi khí Nitơ",
        "q": "Quy trình thổi rửa đường ống đồng (Flushing) sau khi hoàn thiện mối hàn: Loại khí và áp suất thổi xả rác cặn bẩn trong ống là bao nhiêu?",
        "options": [
            "Dùng khí Nitơ khô thổi ngắt quãng với áp suất 0.5 - 0.8 MPa (5-8 bar), dùng giẻ trắng chắn đầu ra kiểm tra độ sạch",
            "Dùng nước máy xịt rửa vào lòng ống đồng",
            "Dùng khí nén máy nén khí ô tô",
            "Thổi bằng hơi thở người"
        ],
        "correct": 0,
        "cad_title": "NITROGEN PIPE FLUSHING & DEBRIS CHECK SETUP SHOP DETAIL",
        "tag": "HVAC-CAD-14"
    },
    {
        "cat": "Hàn ống đồng & Thổi khí Nitơ",
        "q": "Khi hàn ống đồng gần vị trí van tiết lưu điện tử EEV hoặc van dịch vụ dàn lạnh, biện pháp bảo vệ linh kiện khỏi bị cháy hỏng do nhiệt là gì?",
        "options": [
            "Quấn khăn ướt làm mát thân van và xịt nước làm nguội liên tục trong suốt quá trình nung nhiệt hàn",
            "Hàn trực tiếp ngọn lửa rọi thẳng vào thân van",
            "Bọc xốp nylon quanh van",
            "Không cần làm mát"
        ],
        "correct": 0,
        "cad_title": "VALVE OVERHEAT PROTECTION USING WET RAG DURING BRAZING DETAIL",
        "tag": "HVAC-CAD-15"
    },

    # Bọc bảo ôn cách nhiệt (Insulation Lagging) - 4 câu
    {
        "cat": "Bọc bảo ôn cách nhiệt",
        "q": "Nghiệm thu thi công bọc bảo ôn ống đồng (Superlon / Armaflex) cho hệ thống VRV dùng gas R410A/R32: Độ dày bảo ôn ống hơi và ống lỏng tối thiểu quy định trên bản vẽ Shop?",
        "options": [
            "Ống hơi bảo ôn dày 19mm - 25mm, ống lỏng bảo ôn dày 13mm - 19mm; hai ống bọc bảo ôn riêng biệt hoàn toàn",
            "Bọc chung cả ống hơi và ống lỏng vào cùng 1 vỏ bảo ôn",
            "Bảo ôn dày 3mm dán giấy",
            "Không cần bọc bảo ôn ống lỏng"
        ],
        "correct": 0,
        "cad_title": "SEPARATE COPPER PIPE THERMAL INSULATION LAG-STRIP SHOP DETAIL",
        "tag": "HVAC-CAD-16"
    },
    {
        "cat": "Bọc bảo ôn cách nhiệt",
        "q": "Quy trình dán keo mối nối bảo ôn và quấn băng quấn simili ngoài trời: Loại keo chuyên dụng và yêu cầu kỹ thuật dán giáp mối?",
        "options": [
            "Dùng keo dán cao su lưu hóa chuyên dụng (Armaflex 520), quét đều 2 mặt giáp mối, ép chặt không để khe hở; quấn simili đè mí ≥50%",
            "Dùng băng dính hai mặt mỏng dán tạm",
            "Dùng dây thun buộc mối nối",
            "Chỉ lồng bảo ôn vào không dán keo mối nối"
        ],
        "correct": 0,
        "cad_title": "INSULATION ADHESIVE GLUE JOINT & SIMILI WRAP SHOP DETAIL",
        "tag": "HVAC-CAD-17"
    },
    {
        "cat": "Bọc bảo ôn cách nhiệt",
        "q": "Khi ống đồng bảo ôn đi xuyên tường hoặc xuyên sàn bê tông, quy định về ống lồng (Sleeve) và chèn kín chống cháy (Firestop) là gì?",
        "options": [
            "Đặt ống lồng PVC/thép lớn hơn, bảo ôn đi liên tục không bị gián đoạn, chèn kín khe hở bằng keo chống cháy Firestop",
            "Đục bê tông kẹp bẹp ống đồng bảo ôn",
            "Cắt đứt lớp bảo ôn tại vị trí xuyên tường",
            "Trát vữa xi măng trực tiếp đè lên bảo ôn"
        ],
        "correct": 0,
        "cad_title": "WALL & SLAB SLEEVING WITH FIRESTOP PENETRATION SHOP DETAIL",
        "tag": "HVAC-CAD-18"
    },
    {
        "cat": "Bọc bảo ôn cách nhiệt",
        "q": "Nghiệm thu bọc bảo ôn đường ống nước lạnh Chiller đi trên xà gồ mái: Biện pháp chống đọng sương đứt gãy tại vị trí cùm treo (Pipe Hanger Support) là gì?",
        "options": [
            "Sử dụng gối đỡ cùm treo cách nhiệt bằng gỗ cứng ép cao áp hoặc PU Foam đúc sẵn (Insulated Pipe Support Foot)",
            "Kẹp trực tiếp cùm sắt vào lớp bảo ôn xốp mềm",
            "Bỏ qua bảo ôn tại vị trí cùm treo",
            "Trải bao tải đè lên ống"
        ],
        "correct": 0,
        "cad_title": "INSULATED PIPE HANGER SUPPORT WOOD BLOCK SHOP DETAIL",
        "tag": "HVAC-CAD-19"
    },

    # Thử áp suất & Thử kín Nitơ (Pressure Leak Testing) - 5 câu
    {
        "cat": "Thử áp suất & Thử kín Nitơ",
        "q": "Quy trình thử kín áp suất bằng khí Nitơ khô cho đường ống gas VRV R32/R410A gồm 3 giai đoạn áp suất và thời gian giữ áp như thế nào?",
        "options": [
            "Giai đoạn 1: 0.3 MPa (3 phút); Giai đoạn 2: 1.5 MPa (3 phút); Giai đoạn 3: 4.15 MPa (Giữ áp 24 giờ)",
            "Thử 1.0 MPa trong 5 phút",
            "Thử 10.0 MPa trong 10 giây",
            "Thử bằng nước áp lực 0.5 MPa"
        ],
        "correct": 0,
        "cad_title": "3-STAGE NITROGEN PRESSURE LEAK TESTING SCHEMATIC DETAIL",
        "tag": "HVAC-CAD-20"
    },
    {
        "cat": "Thử áp suất & Thử kín Nitơ",
        "q": "Khi nén Nitơ thử áp 4.15 MPa trong 24 giờ, công thức hiệu chỉnh áp suất thử theo sự thay đổi nhiệt độ môi trường là bao nhiêu?",
        "options": [
            "Nhiệt độ môi trường tăng/giảm 1°C thì áp suất trong ống sẽ tăng/giảm tương ứng 0.01 MPa (0.1 bar)",
            "Nhiệt độ thay đổi không ảnh hưởng áp suất",
            "1°C thay đổi áp suất 1.0 MPa",
            "1°C thay đổi áp suất 0.5 bar"
        ],
        "correct": 0,
        "cad_title": "TEMPERATURE PRESSURE CORRECTION CHART FOR TEST ACCEPTANCE",
        "tag": "HVAC-CAD-21"
    },
    {
        "cat": "Thử áp suất & Thử kín Nitơ",
        "q": "Dụng cụ bắt buộc phải lắp tại cụm van nạp Nitơ khi tiến hành thử áp suất hệ thống ống đồng VRV để đảm bảo an toàn lao động là gì?",
        "options": [
            "Van giảm áp Nitơ (Nitrogen Regulator) có 2 đồng hồ đo áp suất cao và áp suất nạp điều chỉnh",
            "Nối trực tiếp bình Nitơ vào hệ thống bằng ống cao su không van giảm áp",
            "Dùng van xả nước inox",
            "Dùng bơm tay xe đạp"
        ],
        "correct": 0,
        "cad_title": "NITROGEN REGULATOR & PRESSURE GAUGES MANIFOLD SHOP DETAIL",
        "tag": "HVAC-CAD-22"
    },
    {
        "cat": "Thử áp suất & Thử kín Nitơ",
        "q": "Phương pháp phát hiện điểm rò rỉ khí Nitơ trên các mối hàn ống đồng và giắc co trong quá trình nén áp suất là gì?",
        "options": [
            "Quét dung dịch xà phòng bọt mịn (Soap Bubble Solution) hoặc máy phát hiện rò siêu âm tại các mối nối",
            "Ngửi bằng mũi",
            "Lắng nghe bằng tai ở khoảng cách 10m",
            "Nhìn màu sơn ống đồng"
        ],
        "correct": 0,
        "cad_title": "SOAP BUBBLE LEAK DETECTION ON BRAZED JOINTS DETAIL",
        "tag": "HVAC-CAD-23"
    },
    {
        "cat": "Thử áp suất & Thử kín Nitơ",
        "q": "Thử áp suất đường ống nước ngưng (Condensate Pipe Test): Phương pháp thử kín tiêu chuẩn trước khi bọc bảo ôn kín là gì?",
        "options": [
            "Bịt các đầu xả, nạp đầy nước vào đường ống duy trì trong 15-30 phút không bị sụt mực nước và không rò rỉ mối nối",
            "Nén hơi 4.0 MPa vào ống nhựa PVC",
            "Chỉ cần nhìn mắt thường không nạp nước",
            "Hút chân không ống nước xả"
        ],
        "correct": 0,
        "cad_title": "GRAVITY DRAINAGE GRAVITY WATER FILL LEAK TEST SHOP DETAIL",
        "tag": "HVAC-CAD-24"
    },

    # Hút chân không & Nạp môi chất lạnh (Evacuation & Refrigerant Charge) - 5 câu
    {
        "cat": "Hút chân không & Nạp môi chất",
        "q": "Quy trình hút chân không hệ thống điều hòa VRV: Yêu cầu đồng hồ đo chân không chuyên dụng (Micron Gauge) và áp suất chân không cần đạt là bao nhiêu?",
        "options": [
            "Sử dụng máy hút chân không 2 cấp, đo bằng Micron Gauge đạt dưới 500 Microns (-755 mmHg), giữ áp chân không trong 1 giờ",
            "Hút bằng máy nén điều hòa cũ trong 10 phút",
            "Hút đạt 5000 Microns là dừng",
            "Chỉ cần mở van xả gas đuổi khí"
        ],
        "correct": 0,
        "cad_title": "2-STAGE VACUUM PUMP & MICRON GAUGE PIPING SETUP SHOP DETAIL",
        "tag": "HVAC-CAD-25"
    },
    {
        "cat": "Hút chân không & Nạp môi chất",
        "q": "Công thức tính toán lượng môi chất lạnh nạp bổ sung (Additional Refrigerant Charge) cho hệ thống VRV căn cứ vào yếu tố nào trên bản vẽ Shop?",
        "options": [
            "Tính theo tổng chiều dài và đường kính của từng đường ống lỏng (Liquid Pipe Length) thi công thực tế",
            "Tính theo chiều dài ống hơi",
            "Nạp tùy thích đến khi dàn nóng mát",
            "Nạp cố định 1kg cho mọi công trình"
        ],
        "correct": 0,
        "cad_title": "LIQUID LINE LENGTH ADDITIONAL REFRIGERANT CALCULATION SHEET",
        "tag": "HVAC-CAD-26"
    },
    {
        "cat": "Hút chân không & Nạp môi chất",
        "q": "Nghiệm thu công tác nạp bổ sung gas R32/R410A vào hệ thống VRV: Thiết bị đo lường bắt buộc phải sử dụng để kiểm soát khối lượng nạp là gì?",
        "options": [
            "Cân điện tử định lượng chính xác (Digital Charging Scale) có độ chia gram",
            "Đo bằng mắt nhìn phin lọc thủy tinh",
            "Ước lượng theo thời gian nạp",
            "Cân bàn quả tạ nông nghiệp"
        ],
        "correct": 0,
        "cad_title": "DIGITAL WEIGHT SCALE LIQUID CHARGING SETUP SHOP DETAIL",
        "tag": "HVAC-CAD-27"
    },
    {
        "cat": "Hút chân không & Nạp môi chất",
        "q": "Khi nạp gas bổ sung ở dạng lỏng vào đường ống lỏng của hệ thống VRV sau khi hút chân không, trạng thái van dịch vụ dàn nóng như thế nào?",
        "options": [
            "Tất cả các van dịch vụ (lỏng và hơi) của dàn nóng vẫn ở trạng thái ĐÓNG KÍN hoàn toàn",
            "Mở hết van dịch vụ dàn nóng rồi mới nạp",
            "Cho dàn nóng chạy nén rồi nạp trực tiếp vào ống hơi",
            "Tháo bỏ ty van xả gas"
        ],
        "correct": 0,
        "cad_title": "LIQUID LINE REFRIGERANT CHARGING VALVE STATUS SHOP DETAIL",
        "tag": "HVAC-CAD-28"
    },
    {
        "cat": "Hút chân không & Nạp môi chất",
        "q": "Nếu hệ thống bị lọt ẩm (Moisture) trong đường ống, chỉ thị trên mắt thăm gas (Sight Glass) sẽ chuyển sang màu gì và cách khắc phục?",
        "options": [
            "Mắt thăm chuyển từ màu XANH (Dry) sang màu VÀNG/HỒNG (Wet); Khắc phục: Thu hồi gas, thay phin lọc rút ẩm và hút chân không lại",
            "Mắt thăm chuyển sang màu đen",
            "Không cần xử lý ẩm",
            "Đổ thêm dầu máy nén vào"
        ],
        "correct": 0,
        "cad_title": "SIGHT GLASS MOISTURE INDICATOR COLOR STATUS SHOP DETAIL",
        "tag": "HVAC-CAD-29"
    },

    # Đi dây điện điều khiển & Truyền thông (Control & Communication Wiring) - 4 câu
    {
        "cat": "Đi dây điện điều khiển ĐHTG",
        "q": "Quy cách cáp điện tín hiệu truyền thông (Communication Cable F1-F2) nối giữa dàn nóng VRV và các dàn lạnh FCU yêu cầu tiêu chuẩn kỹ thuật gì?",
        "options": [
            "Cáp chống nhiễu (Shielded Twisted Pair Cable - 2 core x 0.75-1.25mm²), có lớp vỏ bọc lưới kim loại tiếp địa 1 đầu",
            "Dây điện đơn trần 0.5mm² không bọc vỏ",
            "Dây cáp mạng CAT3 bấm hạt RJ45",
            "Cáp điện lực 3 pha 10mm²"
        ],
        "correct": 0,
        "cad_title": "SHIELDED TWISTED PAIR SIGNAL WIRING BUS DAISY CHAIN DETAIL",
        "tag": "HVAC-CAD-30"
    },
    {
        "cat": "Đi dây điện điều khiển ĐHTG",
        "q": "Sơ đồ đấu nối dây tín hiệu truyền thông giữa cụm dàn nóng và các dàn lạnh trong hệ thống VRV bắt buộc tuân theo cấu trúc nào?",
        "options": [
            "Cấu trúc nối tiếp chuỗi (Daisy Chain Wiring) từ dàn lạnh này sang dàn lạnh tiếp theo, không được nối phân nhánh vòng tròn (Loop/Star)",
            "Đấu nối hình sao khép kín (Star-Loop)",
            "Đấu song song chung cọc nguồn 220V",
            "Đấu tùy tiện không theo thứ tự"
        ],
        "correct": 0,
        "cad_title": "VRV BUS DAISY-CHAIN CONTROL WIRING SCHEMATIC DIAGRAM",
        "tag": "HVAC-CAD-31"
    },
    {
        "cat": "Đi dây điện điều khiển ĐHTG",
        "q": "Khoảng cách đi song song tối thiểu giữa cáp tín hiệu điều khiển VRV và cáp nguồn điện lực 380V để chống nhiễu từ trường là bao nhiêu?",
        "options": [
            "Khoảng cách tối thiểu ≥300mm (hoặc đi trong hai máng cáp/ống thép cách ly riêng biệt)",
            "Đi chung trong cùng 1 ống gen nhựa",
            "Quấn xoắn dây tín hiệu vào dây nguồn 3 pha",
            "Khoảng cách 5mm"
        ],
        "correct": 0,
        "cad_title": "SIGNAL & POWER CABLE SEPARATION CLEARANCE SHOP DETAIL",
        "tag": "HVAC-CAD-32"
    },
    {
        "cat": "Đi dây điện điều khiển ĐHTG",
        "q": "Nghiệm thu đấu nối điều khiển trung tâm BMS (Building Management System) qua cổng Bacnet/Modbus Gateway của hệ thống ĐHTG:",
        "options": [
            "Kiểm tra địa chỉ IP Gateway, cài đặt địa chỉ Central Address cho từng dàn lạnh và test truyền nhận dữ liệu trạng thái/báo lỗi",
            "Chỉ cần cắm nguồn điện 220V vào bộ điều khiển",
            "Dùng công tắc cơ bật tắt trực tiếp",
            "Không cần cài đặt địa chỉ dàn lạnh"
        ],
        "correct": 0,
        "cad_title": "BMS BACNET GATEWAY INTERFACE WIRING SHOP DETAIL",
        "tag": "HVAC-CAD-33"
    },

    # Lắp đặt ống gió & Phụ kiện (Ductwork & Fittings) - 8 câu
    {
        "cat": "Lắp đặt ống gió & Phụ kiện",
        "q": "Nghiệm thu lắp đặt ống gió tôn mạ kẽm (GI Duct): Quy cách độ dày tôn mạ kẽm theo kích thước cạnh lớn nhất của ống gió (TCVN 5687:2010)?",
        "options": [
            "Cạnh <400mm (dày 0.58mm); 400-800mm (dày 0.75mm); 800-1200mm (dày 0.95mm); >1200mm (dày 1.15mm)",
            "Tất cả các cỡ ống gió đều dùng tôn 0.3mm",
            "Cạnh 2000mm dùng tôn 0.5mm",
            "Dùng tôn phẳng không mạ kẽm 0.2mm"
        ],
        "correct": 0,
        "cad_title": "GI GALVANIZED SHEET METAL DUCT THICKNESS & JOINT SPECIFICATION",
        "tag": "HVAC-CAD-34"
    },
    {
        "cat": "Lắp đặt ống gió & Phụ kiện",
        "q": "Kiểm tra kỹ thuật ghép nối ống gió bằng nẹp C / Bích TDC / Bích V: Khoảng cách bắt bu-lông kẹp bích và chèn gioăng làm kín giáp mối là gì?",
        "options": [
            "Dán gioăng cao su/mút tự dán kín mặt bích, khoảng cách kẹp bích/bu-lông góc ≤150mm, bắn silicon kín 4 góc ghép nẹp",
            "Không dán gioăng, bỏ trống góc bích",
            "Bắn đinh rút thưa 500mm 1 con",
            "Dùng dải băng dính quấn quanh bích"
        ],
        "correct": 0,
        "cad_title": "TDC FLANGE & CLEAT DUCT CONNECTION SHOP DETAIL",
        "tag": "HVAC-CAD-35"
    },
    {
        "cat": "Lắp đặt ống gió & Phụ kiện",
        "q": "Nghiệm thu treo ống gió chữ nhật nằm ngang: Khoảng cách tối đa giữa các giá treo (Duct Hangers) và quy cách thanh V đỡ đáy ống?",
        "options": [
            "Khoảng cách giá treo ≤300mm (ống nhỏ) hoặc ≤2500mm; thanh V đỡ đáy tối thiểu V40x40x4mm có lót cao su chống rung",
            "Giá treo cách nhau 10 mét 1 vị trí",
            "Treo ống gió bằng dây thừng treo trần",
            "Không cần thanh V đỡ đáy"
        ],
        "correct": 0,
        "cad_title": "RECTANGULAR DUCT SUPPORT HANGER & ANGLE IRON DETAIL",
        "tag": "HVAC-CAD-36"
    },
    {
        "cat": "Lắp đặt ống gió & Phụ kiện",
        "q": "Yêu cầu kỹ thuật thi công gân tăng cứng (Duct Stiffener) cho đường ống gió chữ nhật có kích thước cạnh lớn từ 800mm trở lên?",
        "options": [
            "Cán gân hình thoi/gân sóng trên thân tôn hoặc gắn thanh V tăng cứng gia cường bên trong/bên ngoài ống gió",
            "Không cần gia cường gân tăng cứng",
            "Dán xốp lên thân ống",
            "Bắn đinh vít dày đặc"
        ],
        "correct": 0,
        "cad_title": "DUCT REINFORCEMENT STIFFENER ANGLE SHOP DETAIL",
        "tag": "HVAC-CAD-37"
    },
    {
        "cat": "Lắp đặt ống gió & Phụ kiện",
        "q": "Nghiệm thu lắp đặt van điều chỉnh lưu lượng gió (VCD - Volume Control Damper) trên đường ống gió:",
        "options": [
            "Lắp tại nhánh rẽ ống gió, tay gạt điều chỉnh dễ thao tác, có vạch chia góc mở và bu-lông hãm cố định vị trí lá van",
            "Bắn chết lá van kín 100% không cho xoay",
            "Giấu van kín trong trần không có tay gạt",
            "Tháo bỏ lá van chỉ giữ vỏ"
        ],
        "correct": 0,
        "cad_title": "MANUAL VOLUME CONTROL DAMPER (VCD) SHOP DETAIL",
        "tag": "HVAC-CAD-38"
    },
    {
        "cat": "Lắp đặt ống gió & Phụ kiện",
        "q": "Khi lắp đặt van chặn lửa (FD - Fire Damper) xuyên tường ngăn cháy: Vị trí lắp đặt van và khoảng cách từ mặt tường đến thân van?",
        "options": [
            "Thân van FD nằm gọn trong bề dày tường ngăn cháy hoặc cách mặt tường ≤150mm, ống gió nối vào van có khớp đứt nhiệt",
            "Lắp van FD cách tường 2 mét",
            "Lắp van FD tự do trên trần nhựa",
            "Bỏ van FD khi đi xuyên tường cháy"
        ],
        "correct": 0,
        "cad_title": "FIRE DAMPER WALL PENETRATION & SLEEVE SHOP DETAIL",
        "tag": "HVAC-CAD-39"
    },
    {
        "cat": "Lắp đặt ống gió & Phụ kiện",
        "q": "Bọc bảo ôn cách nhiệt ống gió tươi / gió lạnh bằng bông thủy tinh (Glasswool) bọc bạc: Yêu cầu tỷ trọng bông và đinh ghim dán giữ bông?",
        "options": [
            "Bông thủy tinh tỷ trọng ≥32kg/m³ (hoặc Rubber Foam), dán đinh ghim đĩa nhôm mật độ 9-12 đinh/m², dán băng keo nhôm kín mối nối",
            "Phủ bông thủy tinh tự do không đinh ghim",
            "Dùng bông tỷ trọng 10kg/m³ bọc lỏng lẻo",
            "Dùng giấy báo bọc quanh ống tôn"
        ],
        "correct": 0,
        "cad_title": "GLASSWOOL FOIL DUCT INSULATION WITH PIN HANGERS DETAIL",
        "tag": "HVAC-CAD-40"
    },
    {
        "cat": "Lắp đặt ống gió & Phụ kiện",
        "q": "Nghiệm thu lắp đặt hộp gió (Plenum Box) và miệng gió nan bầu dục (Linear Grill) trên trần thạch cao:",
        "options": [
            "Hộp gió bọc bảo ôn kín 100%, miệng gió bắt vít chìm hãm chắc chắn vào khung xương trần, mặt miệng gió phẳng khít bề mặt trần",
            "Để hộp gió hở mép xả gió vào trần thạch cao",
            "Lắp miệng gió xệ lệch khỏi trần 20mm",
            "Miệng gió treo lơ lửng bằng dây cước"
        ],
        "correct": 0,
        "cad_title": "CEILING LINEAR SLOT GRILL & PLENUM BOX INSTALLATION DETAIL",
        "tag": "HVAC-CAD-41"
    },

    # Lắp đặt Quạt gió & Thiết bị thông gió (Ventilation Fan Installation) - 5 câu
    {
        "cat": "Lắp đặt Quạt gió thông gió",
        "q": "Nghiệm thu lắp đặt quạt hút khói sự cố / quạt tăng áp ly tâm (Centrifugal Fan) trong phòng kỹ thuật quạt:",
        "options": [
            "Quạt đặt trên lò xo giảm chấn (Vibration Isolator Springs), nối ống gió bằng khớp mềm chịu nhiệt 300°C, vỏ quạt tiếp địa an toàn",
            "Bắt chặt chân quạt trực tiếp xuống sàn không lò xo",
            "Nối cứng ống tôn trực tiếp vào miệng quạt",
            "Quạt treo bằng dây xích xe đạp"
        ],
        "correct": 0,
        "cad_title": "CENTRIFUGAL SMOKE EXTRACTION FAN SPRING ISOLATOR DETAIL",
        "tag": "HVAC-CAD-42"
    },
    {
        "cat": "Lắp đặt Quạt gió thông gió",
        "q": "Quy trình kiểm tra chiều quay động cơ quạt gió 3 pha (3-Phase Fan Motor) trước khi chạy thử tải (Commissioning):",
        "options": [
            "Bật nhấp nháy khởi động (Jogging) quan sát hướng mũi tên chỉ chiều quay trên vỏ quạt; nếu ngược chiều thì đảo 2 trong 3 pha điện",
            "Cho quạt chạy tối đa tốc độ 24 giờ rồi mới kiểm tra",
            "Cứ để quạt quay ngược không ảnh hưởng",
            "Đổi dây trung tính N vào dây pha"
        ],
        "correct": 0,
        "cad_title": "3-PHASE FAN MOTOR ROTATION DIRECTION CHECK DETAIL",
        "tag": "HVAC-CAD-43"
    },
    {
        "cat": "Lắp đặt Quạt gió thông gió",
        "q": "Nghiệm thu quạt thông gió gắn trần nối ống gió (Inline Duct Fan) căn hộ / khu vệ sinh:",
        "options": [
            "Quạt treo độc lập bằng ty ren hãm cao su chống rung, van 1 chiều (Check Damper) ngăn mùi ngược, ống mềm nẹp đai xiết kín",
            "Treo quạt trực tiếp đè lên ống gió",
            "Bỏ van 1 chiều chống mùi",
            "Nối ống mềm bằng dây chun"
        ],
        "correct": 0,
        "cad_title": "INLINE DUCT FAN SUSPENSION & CHECK DAMPER SHOP DETAIL",
        "tag": "HVAC-CAD-44"
    },
    {
        "cat": "Lắp đặt Quạt gió thông gió",
        "q": "Kiểm tra đo đạc độ ồn quạt gió (Noise Level Test) tại vị trí làm việc theo tiêu chuẩn Việt Nam:",
        "options": [
            "Dùng máy đo độ ồn (Sound Level Meter) đo tại khoảng cách 1.5m, độ ồn quạt phòng văn phòng đạt ≤45-55 dBA",
            "Đo độ ồn bằng tai nghe nhạc",
            "Độ ồn 120 dBA vẫn đạt chuẩn",
            "Không cần đo độ ồn"
        ],
        "correct": 0,
        "cad_title": "FAN ACOUSTIC SOUND LEVEL MEASUREMENT LOCATION SHOP DETAIL",
        "tag": "HVAC-CAD-45"
    },
    {
        "cat": "Lắp đặt Quạt gió thông gió",
        "q": "Thiết bị giảm âm quạt gió (Silencer / Sound Attenuator) lắp trên đường ống gió quạt hút có cấu tạo và chức năng gì?",
        "options": [
            "Vỏ tôn bọc đệm vật liệu hấp thụ âm thanh (Bông khoáng / Bông thủy tinh đục lỗ), giảm độ ồn lan truyền theo dòng không khí",
            "Là van đóng mở dòng khí",
            "Là lọc bụi thô",
            "Là bộ gia nhiệt điện"
        ],
        "correct": 0,
        "cad_title": "DUCT SOUND ATTENUATOR / SILENCER SHOP DRAWING DETAIL",
        "tag": "HVAC-CAD-46"
    },

    # Thử kín ống gió & Nạp khí (Duct Leakage Testing DW/144) - 4 câu
    {
        "cat": "Thử kín ống gió SMACNA/DW144",
        "q": "Nghiệm thu thử kín đường ống gió áp suất cao/trung bình theo tiêu chuẩn DW/144 hoặc SMACNA: Phương pháp thử kín bằng quạt thử áp và đĩa lỗ (Orifice Plate)?",
        "options": [
            "Bịt kín hai đầu đoạn ống thử, bơm khí tạo áp suất thiết kế, đo lưu lượng khí rò rỉ qua đĩa lỗ orifice không vượt quá giới hạn cho phép",
            "Thử kín bằng cách đổ nước vào ống gió",
            "Nhìn mắt thường tìm khe hở",
            "Thử kín bằng khói thuốc lá"
        ],
        "correct": 0,
        "cad_title": "DW144 DUCT LEAKAGE TESTING RIG & ORIFICE METER DETAIL",
        "tag": "HVAC-CAD-47"
    },
    {
        "cat": "Thử kín ống gió SMACNA/DW144",
        "q": "Phương pháp thử kín ống gió bằng máy tạo khói (Smoke Test Method) áp dụng tại công trường:",
        "options": [
            "Bơm khói đặc vào đoạn ống gió đã bịt kín và tạo áp nhẹ, quan sát các vị trí khớp nối nẹp bích/mối ghép xem có khói rò rỉ ra ngoài",
            "Đốt rơm tạo khói trong ống gió",
            "Xịt nước xà phòng lên toàn bộ thân tôn",
            "Thử khói khi quạt gió đang chạy tối đa"
        ],
        "correct": 0,
        "cad_title": "SMOKE GENERATOR DUCT LEAKAGE TEST SETUP SHOP DETAIL",
        "tag": "HVAC-CAD-48"
    },
    {
        "cat": "Thử kín ống gió SMACNA/DW144",
        "q": "Tỷ lệ phần trăm (%) tổng độ dài đường ống gió bắt buộc phải chọn ngẫu nhiên để thử kín áp suất tại công trường theo quy định nghiệm thu:",
        "options": [
            "Tối thiểu 10% - 20% tổng diện tích/độ dài ống gió cho hệ thống áp suất trung bình/cao (hoặc 100% cho tuyến ống nguy hiểm/hút khói)",
            "Chỉ thử 0.1% độ dài",
            "Không bắt buộc thử kín ống gió",
            "Chỉ thử đoạn ống ngắn 0.5m"
        ],
        "correct": 0,
        "cad_title": "RANDOM DUCT SECTION SELECTION FOR SMACNA TESTING DETAIL",
        "tag": "HVAC-CAD-49"
    },
    {
        "cat": "Thử kín ống gió SMACNA/DW144",
        "q": "Nghiệm thu cân bằng khí (TAB - Testing, Adjusting, Balancing) toàn bộ hệ thống ĐHTG trước khi bàn giao đưa vào sử dụng:",
        "options": [
            "Đo lưu lượng gió tại tất cả các miệng thổi/hồi bằng Anemometer/Balancing Hood, chỉnh van VCD đạt sai số lưu lượng nằm trong khoảng ±10%",
            "Chỉ cần kiểm tra quạt có quay hay không",
            "Chỉnh van VCD mở tự do 100%",
            "Sai số lưu lượng ±50% vẫn nghiệm thu"
        ],
        "correct": 0,
        "cad_title": "TAB AIRFLOW BALANCING HOOD MEASUREMENT SHOP DETAIL",
        "tag": "HVAC-CAD-50"
    }
]

print(f"Generated {len(hvac_tasks)} detailed practical HVAC questions.")

# ---------------------------------------------------------------------
# 2. GENERATE 50 HIGH-QUALITY CAD SHOP DRAWING SVG BLUEPRINTS
# ---------------------------------------------------------------------
data_dir = r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\data'
os.makedirs(data_dir, exist_ok=True)

for idx, task in enumerate(hvac_tasks, start=1):
    cad_title = task["cad_title"]
    cad_tag = task["tag"]
    cat_name = task["cat"]
    
    svg_content = f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1000 650" width="100%" height="100%">
  <!-- CAD Shop Drawing Dark Engineering Background -->
  <rect width="1000" height="650" fill="#0f172a"/>
  
  <!-- Technical CAD Grid Lines -->
  <defs>
    <pattern id="cadGrid" width="50" height="50" patternUnits="userSpaceOnUse">
      <path d="M 50 0 L 0 0 0 50" fill="none" stroke="#1e293b" stroke-width="1"/>
      <path d="M 10 0 L 10 50 M 20 0 L 20 50 M 30 0 L 30 50 M 40 0 L 40 50 M 0 10 L 50 10 M 0 20 L 50 20 M 0 30 L 50 30 M 0 40 L 50 40" fill="none" stroke="#111827" stroke-width="0.5"/>
    </pattern>
    <linearGradient id="cyanGrad" x1="0%" y1="0%" x2="100%" y2="100%">
      <stop offset="0%" stop-color="#38bdf8"/>
      <stop offset="100%" stop-color="#0284c7"/>
    </linearGradient>
    <linearGradient id="amberGrad" x1="0%" y1="0%" x2="100%" y2="100%">
      <stop offset="0%" stop-color="#fbbf24"/>
      <stop offset="100%" stop-color="#d97706"/>
    </linearGradient>
  </defs>
  <rect width="1000" height="650" fill="url(#cadGrid)"/>

  <!-- Outer Border Frame -->
  <rect x="15" y="15" width="970" height="620" fill="none" stroke="#3b82f6" stroke-width="2.5"/>
  <rect x="22" y="22" width="956" height="606" fill="none" stroke="#1e3a8a" stroke-width="1"/>

  <!-- CAD Drawing Main Header Banner -->
  <rect x="22" y="22" width="956" height="45" fill="#1e293b"/>
  <text x="35" y="50" fill="#38bdf8" font-family="monospace, Arial" font-size="16" font-weight="bold">VINCONS MEP CAD SHOP DRAWING — {cat_name.upper()}</text>
  <text x="780" y="50" fill="#fbbf24" font-family="monospace, Arial" font-size="14" font-weight="bold">DWG NO: {cad_tag}</text>
  <line x1="22" y1="67" x2="978" y2="67" stroke="#3b82f6" stroke-width="1.5"/>

  <!-- CAD Blueprint Graphics Canvas Area -->
  <g transform="translate(60, 90)">
    <!-- Equipment / Pipe Assembly Graphic Shapes -->
    <rect x="80" y="40" width="380" height="260" fill="#1e293b" stroke="#38bdf8" stroke-width="3" rx="6"/>
    <rect x="100" y="60" width="160" height="220" fill="#0f172a" stroke="#0284c7" stroke-width="2" rx="4"/>
    <circle cx="180" cy="170" r="55" fill="none" stroke="#fbbf24" stroke-width="3.5" stroke-dasharray="8 4"/>
    <circle cx="180" cy="170" r="18" fill="#d97706"/>
    <line x1="180" y1="115" x2="180" y2="225" stroke="#fbbf24" stroke-width="2"/>
    <line x1="125" y1="170" x2="235" y2="170" stroke="#fbbf24" stroke-width="2"/>

    <!-- Piping / Ductwork Lines -->
    <path d="M 260 110 L 520 110 L 520 280 L 680 280" fill="none" stroke="#4ade80" stroke-width="5" stroke-linecap="round"/>
    <path d="M 260 140 L 490 140 L 490 310 L 680 310" fill="none" stroke="#f43f5e" stroke-width="4" stroke-linecap="round"/>

    <!-- Valve / Sensor Fittings Icons -->
    <polygon points="510,110 530,100 530,120" fill="#4ade80"/>
    <polygon points="530,110 510,100 510,120" fill="#4ade80"/>
    <circle cx="520" cy="110" r="6" fill="#fbbf24"/>

    <!-- Insulation Layer Coating Representation -->
    <rect x="270" y="102" width="200" height="16" fill="none" stroke="#a855f7" stroke-width="2" stroke-dasharray="4 2"/>
    <text x="310" y="95" fill="#a855f7" font-family="monospace" font-size="11">SUPERLON INSULATION 25mm</text>

    <!-- Engineering Dimensioning Lines -->
    <line x1="80" y1="320" x2="460" y2="320" stroke="#64748b" stroke-width="1.5"/>
    <line x1="80" y1="312" x2="80" y2="328" stroke="#64748b" stroke-width="1.5"/>
    <line x1="460" y1="312" x2="460" y2="328" stroke="#64748b" stroke-width="1.5"/>
    <text x="240" y="340" fill="#94a3b8" font-family="monospace" font-size="12" text-anchor="middle">CLEARANCE ≥ 1500 mm</text>

    <line x1="40" y1="40" x2="40" y2="300" stroke="#64748b" stroke-width="1.5"/>
    <line x1="32" y1="40" x2="48" y2="40" stroke="#64748b" stroke-width="1.5"/>
    <line x1="32" y1="300" x2="48" y2="300" stroke="#64748b" stroke-width="1.5"/>
    <text x="25" y="175" fill="#94a3b8" font-family="monospace" font-size="12" transform="rotate(-90 25 175)" text-anchor="middle">HEIGHT = 1200 mm</text>

    <!-- Technical Callout Leader Notes -->
    <path d="M 460 70 L 580 30 L 660 30" fill="none" stroke="#38bdf8" stroke-width="1.5"/>
    <circle cx="460" cy="70" r="4" fill="#38bdf8"/>
    <text x="670" y="34" fill="#38bdf8" font-family="monospace" font-size="12" font-weight="bold">ITEM 01: VRV OUTDOOR / FCU UNIT</text>

    <path d="M 520 280 L 600 360 L 660 360" fill="none" stroke="#4ade80" stroke-width="1.5"/>
    <circle cx="520" cy="280" r="4" fill="#4ade80"/>
    <text x="670" y="364" fill="#4ade80" font-family="monospace" font-size="12" font-weight="bold">ITEM 02: REFRIGERANT / DUCT LINE ({task["tag"]})</text>
  </g>

  <!-- Title Block Standard CAD Form (Góc Phải Bên Dưới) -->
  <g transform="translate(560, 480)">
    <rect x="0" y="0" width="418" height="148" fill="#1e293b" stroke="#3b82f6" stroke-width="2"/>
    <line x1="0" y1="30" x2="418" y2="30" stroke="#334155" stroke-width="1"/>
    <line x1="0" y1="60" x2="418" y2="60" stroke="#334155" stroke-width="1"/>
    <line x1="0" y1="95" x2="418" y2="95" stroke="#334155" stroke-width="1"/>
    <line x1="200" y1="0" x2="200" y2="95" stroke="#334155" stroke-width="1"/>

    <text x="10" y="20" fill="#94a3b8" font-family="Arial" font-size="10">DỰ ÁN / PROJECT:</text>
    <text x="10" y="48" fill="#f8fafc" font-family="Arial" font-size="11" font-weight="bold">BẾN DU THUYỀN VŨ YÊN HẢI PHÒNG</text>
    <text x="210" y="20" fill="#94a3b8" font-family="Arial" font-size="10">HẠNG MỤC / PACKAGE:</text>
    <text x="210" y="48" fill="#38bdf8" font-family="Arial" font-size="11" font-weight="bold">ĐIỀU HOÀ THÔNG GIÓ (HVAC)</text>

    <text x="10" y="76" fill="#94a3b8" font-family="Arial" font-size="10">BẢN VẼ / DRAWING TITLE:</text>
    <text x="10" y="89" fill="#fbbf24" font-family="Arial" font-size="10" font-weight="bold">{cad_title[:45]}</text>
    
    <text x="210" y="76" fill="#94a3b8" font-family="Arial" font-size="10">NGHIỆM THU / FIELD TASK:</text>
    <text x="210" y="89" fill="#4ade80" font-family="Arial" font-size="10" font-weight="bold">{cat_name}</text>

    <text x="10" y="115" fill="#cbd5e1" font-family="monospace" font-size="10">ĐƠN VỊ THI CÔNG: VINCONS MEP CONTRACTOR</text>
    <text x="10" y="135" fill="#94a3b8" font-family="monospace" font-size="10">TỶ LỆ: 1:25 | KHỔ A2 | NGÀY: 26/08/2026 | BẢN VẼ SHOP KHỔ A2</text>
  </g>
</svg>'''

    svg_filename = f"hvac_cad_{idx}.svg"
    svg_filepath = os.path.join(data_dir, svg_filename)
    with open(svg_filepath, 'w', encoding='utf-8') as f:
        f.write(svg_content)
    
    # Update task image field
    task["image"] = f"data/{svg_filename}"

print("✅ Generated 50 SVG CAD Shop Drawing Blueprints in data/ directory.")

# ---------------------------------------------------------------------
# 3. SAVE 50 PRACTICAL HVAC ESSAYS TO QUESTIONS.JS
# ---------------------------------------------------------------------
questions_db_path = r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\questions.js'
with open(questions_db_path, 'r', encoding='utf-8') as f:
    text = f.read()

questions = json.loads(text[text.find('['):text.rfind(']')+1])

# Remove old practical HVAC if any
questions = [q for q in questions if not (q.get('category') == 'Thực hành ĐHTG' or q.get('exam_set') == 'Tự luận - Thực hành ĐHTG')]

formatted_hvac_questions = []
for idx, task in enumerate(hvac_tasks, start=1):
    q_obj = {
        "id": f"q_hvac_prac_{idx}",
        "type": "multiple_choice",
        "category": "Thực hành ĐHTG",
        "exam_set": "Tự luận - Thực hành ĐHTG",
        "question": task["q"],
        "options": task["options"],
        "correct_index": task["correct"],
        "image": task["image"]
    }
    formatted_hvac_questions.append(q_obj)

questions.extend(formatted_hvac_questions)
print(f"Total questions in bank: {len(questions)} (including {len(formatted_hvac_questions)} Practical HVAC questions).")

new_js = f"// File questions.js - Tích hợp đầy đủ 50 câu hỏi Tự luận - Thực hành ĐHTG đi kèm Bản vẽ CAD Shop Drawing\nconst QUESTIONS = {json.dumps(questions, ensure_ascii=False, indent=2)};\n"

js_paths = [
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\questions.js',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\apk_unpacked\assets\www\questions.js',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\android_build\android\app\src\main\assets\public\questions.js',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\android_build\www\questions.js'
]

for p in js_paths:
    if os.path.exists(os.path.dirname(p)):
        with open(p, 'w', encoding='utf-8') as f:
            f.write(new_js)
        print(f"✅ Saved questions.js to {p}")

# Also copy data/ SVG files to target build dirs
target_dirs = [
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\apk_unpacked\assets\www\data',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\android_build\android\app\src\main\assets\public\data',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\android_build\www\data'
]

for tdir in target_dirs:
    os.makedirs(tdir, exist_ok=True)
    for idx in range(1, 51):
        sname = f"hvac_cad_{idx}.svg"
        src_svg = os.path.join(data_dir, sname)
        if os.path.exists(src_svg):
            shutil.copy(src_svg, os.path.join(tdir, sname))

# ---------------------------------------------------------------------
# 4. EXPORT EXCEL & WORD BACKUP FILES TO USER'S DESTINATION DIRECTORY
# ---------------------------------------------------------------------
dest_folder = r'D:\LINH TINH\AI\Khung chương trình đào tạo nghề Điện - CTN\De thi Ý gửi\phần đáp án'
os.makedirs(dest_folder, exist_ok=True)

excel_out = os.path.join(dest_folder, 'Bo_De_Thi_Tu_Luan_Thuc_Hanh_DHTG_50_Cau_Kem_Ban_Ve_CAD.xlsx')
word_out = os.path.join(dest_folder, 'Bo_De_Thi_Tu_Luan_Thuc_Hanh_DHTG_50_Cau_Kem_Ban_Ve_CAD.docx')

# Create Excel Backup
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Đề thi Thực hành ĐHTG 50 câu"

headers = ["STT", "Mã Cụm Công Việc", "Danh Mục Nghiệm Thu", "Nội Dung Câu Hỏi Tự Luận - Thực Hành", "Đáp Án A", "Đáp Án B", "Đáp Án C", "Đáp Án D", "Đáp Án Đúng", "File Bản Vẽ CAD đính kèm"]
ws.append(headers)

# Styling
header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")

for col_num, header in enumerate(headers, 1):
    cell = ws.cell(row=1, column=col_num)
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

for idx, q in enumerate(formatted_hvac_questions, 1):
    row = [
        idx,
        q["id"],
        q["category"],
        q["question"],
        q["options"][0],
        q["options"][1],
        q["options"][2],
        q["options"][3],
        chr(65 + q["correct_index"]),
        q["image"]
    ]
    ws.append(row)

ws.column_dimensions['A'].width = 6
ws.column_dimensions['B'].width = 16
ws.column_dimensions['C'].width = 25
ws.column_dimensions['D'].width = 60
ws.column_dimensions['E'].width = 30
ws.column_dimensions['F'].width = 30
ws.column_dimensions['G'].width = 30
ws.column_dimensions['H'].width = 30
ws.column_dimensions['I'].width = 12
ws.column_dimensions['J'].width = 25

wb.save(excel_out)
print(f"✅ Created Excel backup: {excel_out}")

# Create Word Backup
doc = docx.Document()
doc.add_heading('BỘ ĐỀ THI TỰ LUẬN - THỰC HÀNH HỆ THỐNG ĐIỀU HÒA THÔNG GIÓ (50 CÂU)', level=0)
p_sub = doc.add_paragraph('Chuyên ngành: Giám sát Thi công & Nghiệm thu Hệ thống ĐHTG - Công trình Vũ Yên Hải Phòng')
p_sub.runs[0].font.italic = True

for idx, q in enumerate(formatted_hvac_questions, 1):
    h = doc.add_heading(f"Câu {idx}: [{q['category']}] (Mã CAD: {q['image']})", level=2)
    doc.add_paragraph(q["question"])
    
    for opt_idx, opt in enumerate(q["options"]):
        prefix = f"  {chr(65 + opt_idx)}. "
        p_opt = doc.add_paragraph(prefix + opt)
        if opt_idx == q["correct_index"]:
            p_opt.runs[0].font.bold = True
            p_opt.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph(f"👉 ĐÁP ÁN ĐÚNG: {chr(65 + q['correct_index'])}")
    doc.add_paragraph("-" * 50)

doc.save(word_out)
print(f"✅ Created Word backup: {word_out}")
