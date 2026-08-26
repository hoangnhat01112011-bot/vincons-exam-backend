import json
import os
import shutil
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
import docx
from docx.shared import Pt, RGBColor

# 1. 50 PROFESSIONAL PCCC THEORY QUESTIONS ACCORDING TO TCVN & QCVN STANDARDS
pccc_questions = [
    # 1. Hệ thống báo cháy tự động (TCVN 5738:2021) - 10 câu
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Theo TCVN 5738:2021, khoảng cách tối đa từ bất kỳ điểm nào trên trần nhà đến đầu báo khói tự động trong phòng phẳng có chiều cao dưới 3.5m là bao nhiêu?",
        "opts": ["Bán kính bảo vệ r = 6.5m (diện tích bảo vệ đến 85m²)", "Bán kính bảo vệ r = 10m", "Bán kính bảo vệ r = 3m", "Bán kính bảo vệ r = 15m"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Khoảng cách tối thiểu từ đầu báo cháy tự động (khói/nhiệt) đến mép tường nhà hoặc góc tường theo tiêu chuẩn TCVN 5738:2021 là bao nhiêu?",
        "opts": ["Tối thiểu ≥ 0.1m (10cm)", "Tối thiểu ≥ 1.5m", "Phải lắp sát góc tường 0cm", "Tối thiểu ≥ 3.0m"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Khoảng cách lắp đặt tối thiểu từ đầu báo cháy tự động đến miệng thổi của hệ thống thông gió/điều hòa không khí để tránh dòng gió làm lệch khói là bao nhiêu?",
        "opts": ["Tối thiểu ≥ 1.5m", "Tối thiểu ≥ 0.2m", "Có thể lắp trực tiếp sát miệng thổi 0.1m", "Tối thiểu ≥ 5.0m"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Độ cao lắp đặt tiêu chuẩn cho Nút ấn báo cháy bằng tay (Manual Call Point) tính từ mặt sàn hoàn thiện đến tâm nút ấn là bao nhiêu?",
        "opts": ["Độ cao từ 1.3m đến 1.5m", "Độ cao từ 0.5m đến 0.8m", "Độ cao từ 2.0m đến 2.5m", "Độ cao sát mặt sàn 0.2m"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Theo TCVN 5738:2021, dây tín hiệu báo cháy (Loop/Zone) khi đi xuyên qua tường hoặc sàn nhà bắt buộc phải thi công kỹ thuật như thế nào?",
        "opts": ["Phải lồng trong ống bảo vệ (PVC/thép) và chèn kín khe hở bằng vật liệu chống cháy", "Đi dây trần trực tiếp qua lỗ đục gạch", "Dùng băng dính quấn quanh dây", "Kẹp đinh thép trực tiếp vào bê tông"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Yêu cầu nguồn điện dự phòng (Acquy) cấp cho Trung tâm báo cháy tự động phải duy trì hoạt động tối thiểu bao nhiêu giờ ở chế độ thường trực và chế độ báo cháy?",
        "opts": ["Duy trì tối thiểu 24 giờ ở chế độ thường trực và 1 giờ ở chế độ báo cháy liên tục", "Duy trì 1 giờ thường trực và 5 phút báo cháy", "Duy trì 12 giờ thường trực và 15 phút báo cháy", "Không cần acquy dự phòng"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Đầu báo cháy nhiệt gia tăng (Rate-of-Rise Heat Detector) kích hoạt phát tín hiệu báo cháy dựa trên nguyên lý nào?",
        "opts": ["Phát hiện tốc độ gia tăng nhiệt độ môi trường vượt quá ngưỡng quy định (thường từ 8°C - 10°C/phút)", "Phát hiện nồng độ khói mờ trong không khí", "Phát hiện tia cực tím từ ngọn lửa", "Phát hiện rò rỉ khí gas"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC BCCC Bậc 2 & Bậc 3",
        "q": "Điện áp hoạt động tiêu chuẩn của các thiết bị đầu báo cháy, chuông, còi và đèn chớp trong hệ thống báo cháy tự động kênh/địa chỉ thường là bao nhiêu?",
        "opts": ["24V DC (Điện áp một chiều 24V an toàn)", "220V AC", "380V AC", "12V AC"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Khoảng cách giữa các Nút ấn báo cháy bằng tay bố trí dọc theo hành lang thoát nạn không được vượt quá bao nhiêu mét?",
        "opts": ["Không vượt quá 50m (hoặc 30m đối với khu vực nguy hiểm cháy cao)", "Không vượt quá 100m", "Không vượt quá 150m", "Chỉ lắp 1 nút duy nhất ở cổng chính"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Chuông báo cháy và còi/đèn chớp báo cháy lắp đặt tại khu vực hành lang nhà phải đảm bảo độ ồn tối thiểu cao hơn độ ồn môi trường xung quanh là bao nhiêu dBA?",
        "opts": ["Cao hơn độ ồn môi trường ít nhất 15 dBA (hoặc đạt tối thiểu 75 dBA tại khoảng cách 3m)", "Cao hơn 2 dBA", "Đạt tối đa 30 dBA", "Không quy định độ ồn"],
        "correct": 0
    },

    # 2. Hệ thống chữa cháy tự động Sprinkler/Deluge (TCVN 7336:2021) - 10 câu
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Theo TCVN 7336:2021, nhiệt độ kích hoạt nổ chốt thủy tinh màu ĐỎ của đầu phun Sprinkler tiêu chuẩn (68°C) áp dụng cho môi trường làm việc nào?",
        "opts": ["Môi trường nhiệt độ thường, nhiệt độ môi trường tối đa không vượt quá 38°C", "Khu vực lò hơi 150°C", "Khu vực đông lạnh -20°C", "Khu vực sấy nông sản 100°C"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Khoảng cách từ tấm định hướng của đầu phun Sprinkler hướng xuống (Pendent) đến trần nhà phẳng tối thiểu và tối đa là bao nhiêu?",
        "opts": ["Từ 0.075m (7.5cm) đến 0.15m (15cm) [hoặc tối đa 0.3m]", "Sát trần 0cm", "Từ 1.0m đến 2.0m", "Từ 0.5m đến 0.8m"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Bán kính bảo vệ hoặc khoảng cách tối đa giữa 2 đầu phun Sprinkler trên cùng một nhánh ống chữa cháy trong không gian nguy hiểm cháy trung bình (Nhóm II) là bao nhiêu?",
        "opts": ["Khoảng cách giữa 2 đầu phun ≤ 3.5m đến 4.0m (bán kính bảo vệ r ≈ 2.0m - 2.1m)", "Khoảng cách ≤ 8.0m", "Khoảng cách ≤ 10.0m", "Khoảng cách 1.0m"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Áp suất làm việc tối thiểu tại đầu phun Sprinkler xa nhất và cao nhất của hệ thống chữa cháy Sprinkler tự động phải đạt bao nhiêu bar (MPa)?",
        "opts": ["Tối thiểu ≥ 0.5 bar (0.05 MPa / 0.5 kg/cm²)", "Tối thiểu ≥ 10.0 bar", "Tối thiểu ≥ 0.01 bar", "Tối thiểu ≥ 20.0 bar"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Cụm van báo động (Alarm Valve / Alarm Check Valve) trong hệ thống Sprinkler đường ống ướt có chức năng chính là gì?",
        "opts": ["Mở cho nước chảy đến đầu phun khi có cháy, đồng thời kích hoạt chuông nước (Water Motor Gong) và công tắc áp lực báo về tủ trung tâm", "Khóa chặt nước không cho nước chảy", "Giảm áp suất nước xuống 0", "Hút khí vào đường ống"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Trường hợp khu vực thi công có trần treo thạch cao, quy định lắp đặt đầu phun Sprinkler như thế nào?",
        "opts": ["Lắp đầu phun Sprinkler hướng xuống (Pendent) nhô ra dưới trần thạch cao; nếu khoảng không gian trên trần >0.75m có vật liệu cháy thì phải lắp thêm Sprinkler quay lên trên trần", "Chỉ lắp 1 đầu trên trần bê tông giấu kín", "Không được lắp Sprinkler dưới trần thạch cao", "Cắt thủng trần thạch cao 1m²"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Hệ thống chữa cháy tràn ngập (Deluge System) sử dụng loại đầu phun nào và nguyên lý kích hoạt xả nước ra sao?",
        "opts": ["Sử dụng đầu phun hở (Open Sprinkler); khi trung tâm báo cháy nhận tín hiệu từ 2 đầu báo cháy sẽ mở van Deluge Valve xả nước đồng loạt qua tất cả các đầu phun", "Sử dụng đầu phun kín chốt thủy tinh 68°C", "Mở từng đầu phun thủ công bằng tay", "Xả bằng bọt foam"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Đường kính định danh (DN) tối thiểu của đường ống cấp nước chữa cháy chính dẫn tới cụm van Alarm Valve theo tiêu chuẩn TCVN 7336:2021 là bao nhiêu?",
        "opts": ["Tối thiểu DN80 hoặc DN100 (tùy thuộc vào số lượng đầu phun)", "DN15", "DN25", "DN32"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Quy định về thử kín và thử áp lực thủy lực đối với mạng đường ống chữa cháy Sprinkler trước khi bọc cách nhiệt hoặc đưa vào nghiệm thu là bao nhiêu?",
        "opts": ["Thử áp lực bằng nước với áp suất bằng 1.5 lần áp suất làm việc (tối thiểu 1.0 MPa đến 1.5 MPa) duy trì trong 2 giờ không sụt áp", "Thử áp suất 0.1 MPa trong 5 phút", "Thử nén hơi 0.05 MPa", "Không cần thử áp lực"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Công tắc dòng chảy (Flow Switch) được lắp đặt trên các nhánh ống chữa cháy Sprinkler tầng nhằm mục đích gì?",
        "opts": ["Phát hiện dòng nước di chuyển khi có đầu phun Sprinkler bị nổ và gửi tín hiệu báo chính xác vị trí tầng đang có cháy về tủ trung tâm", "Khóa đường ống nước khi có sự cố", "Tăng áp suất nước trong ống", "Lọc rác trong đường ống"],
        "correct": 0
    },

    # 3. Hệ thống họng nước chữa cháy & Trụ nước ngoài nhà (TCVN 3890:2023) - 8 câu
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Chiều cao lắp đặt tiêu chuẩn của Hộp chữa cháy vách tường (chứa van họng nước D50/D65, cuộn vòi và lăng phun) tính từ mặt sàn đến tâm van họng nước là bao nhiêu?",
        "opts": ["Độ cao 1.25m (±0.05m)", "Độ cao 2.5m", "Độ cao 0.3m", "Đặt nằm trên sàn nhà"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Áp suất tự do tối thiểu tại miệng lăng phun chữa cháy vách tường khi hệ thống hoạt động xả nước để đảm bảo chiều cao cột nước chữa cháy đạt hiệu quả là bao nhiêu?",
        "opts": ["Tối thiểu ≥ 2.0 bar (0.2 MPa / cột nước chữa cháy ≥ 6m - 10m)", "Tối thiểu 0.1 bar", "Tối thiểu 20.0 bar", "Tối thiểu 50.0 bar"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Độ dài tiêu chuẩn của một cuộn vòi chữa cháy dải vải bọc cao su (D50 hoặc D65) trang bị trong tủ PCCC vách tường là bao nhiêu mét?",
        "opts": ["Độ dài 20m (hoặc 30m theo thiết kế)", "Độ dài 5m", "Độ dài 100m", "Độ dài 50m"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Trụ nước chữa cháy ngoài nhà 3 cửa (1 cửa D125/D100 và 2 cửa D65) phải được lắp đặt cách mép đường giao thông tối đa bao nhiêu mét?",
        "opts": ["Cách mép đường không quá 2.5m và cách tường nhà tối thiểu 5.0m", "Cách mép đường 20m", "Lắp giữa lòng đường giao thông", "Cách tường nhà 0.1m"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Khoảng cách tối đa giữa các Trụ nước chữa cháy ngoài nhà bố trí dọc theo đường giao thông nội bộ công trình là bao nhiêu mét?",
        "opts": ["Khoảng cách giữa 2 trụ nước không vượt quá 150m", "Khoảng cách không vượt quá 500m", "Khoảng cách không vượt quá 30m", "Chỉ lắp 1 trụ cho toàn khu"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Họng tiếp nước chữa cháy ngoài nhà (Siamese Connection D65x2) dành cho xe chữa cháy bơm cấp nước vào tòa nhà bắt buộc trang bị van một chiều nhằm mục đích gì?",
        "opts": ["Cho phép xe chữa cháy bơm nước một chiều vào mạng đường ống tòa nhà và ngăn nước trong nhà chảy ngược ra ngoài", "Cho nước chảy tự do hai chiều", "Xả bớt áp suất nước ngoài đường", "Hút khí vào tòa nhà"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Vật liệu chế tạo lăng phun và khớp nối cuộn vòi chữa cháy vách tường D50/D65 tiêu chuẩn PCCC Việt Nam thường là gì?",
        "opts": ["Hợp kim nhôm đúc hoặc đồng thau mạ niken chịu lực va đập và chống ăn mòn", "Nhựa tái chế mỏng", "Thủy tinh mỏng", "Gỗ nén ép"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Quy trình nghiệm thu kiểm tra thử cuộn vòi chữa cháy trước khi đưa vào trang bị sử dụng:",
        "opts": ["Thử áp lực thủy lực cuộn vòi ở áp suất thử 1.2 - 1.6 MPa trong 2 phút không phồng rách, khớp nối xoay nhẹ khít gioăng", "Chỉ cần mở vòi ra phơi nắng", "Thổi hơi bằng miệng", "Ngâm vòi trong xăng"],
        "correct": 0
    },

    # 4. Cụm bơm chữa cháy & Tủ điều khiển (TCVN 02:2020/BCA) - 8 câu
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Theo QCVN 02:2020/BCA, cụm bơm chữa cháy chính của tòa nhà bắt buộc phải trang bị các loại bơm nào?",
        "opts": ["Bơm chữa cháy động cơ điện chính + Bơm dự phòng (bơm động cơ Diesel hoặc điện nguồn riêng) + Bơm bù áp Jockey", "Chỉ cần 1 máy bơm nước gia đình 1HP", "Không cần bơm dự phòng", "Chỉ dùng 1 máy bơm tay"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Chức năng chính của Máy bơm bù áp (Jockey Pump) trong cụm máy bơm chữa cháy tự động là gì?",
        "opts": ["Bù đắp lượng nước rò rỉ nhỏ để duy trì áp suất thường trực trong đường ống ở mức thiết kế mà không làm khởi động bơm chính", "Dùng để chữa cháy chính khi có đám cháy lớn", "Hút kiệt nước trong bể", "Đổi chiều dòng chảy"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Cơ chế tự động khởi động máy bơm chữa cháy chính khi áp suất đường ống bị sụt giảm do nổ Sprinkler hoặc mở họng nước là nhờ thiết bị nào?",
        "opts": ["Công tắc áp lực (Pressure Switch) lắp trên bình tích áp / cụm ống góp máy bơm", "Bật công tắc bằng tay", "Công tắc hành trình cửa", "Cảm biến nhiệt độ không khí"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Theo quy chuẩn PCCC, sau khi máy bơm chữa cháy điện chính hoặc bơm Diesel tự động khởi động chạy chữa cháy, quy định dừng máy bơm như thế nào?",
        "opts": ["Máy bơm chữa cháy chính KHÔNG ĐƯỢC TỰ ĐỘNG DỪNG, chỉ được dừng máy thủ công bằng tay tại tủ điều khiển sau khi đã dập tắt cháy", "Tự động dừng sau 1 phút", "Tự động dừng khi áp suất tăng cao", "Dừng khi hết nước bể"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Yêu cầu dung tích bình nhiên liệu Dầu Diesel cấp cho máy bơm chữa cháy động cơ Diesel dự phòng phải đảm bảo máy chạy liên tục tối thiểu bao nhiêu giờ?",
        "opts": ["Đảm bảo cho máy bơm Diesel vận hành liên tục 100% tải trong thời gian tối thiểu từ 3 đến 4 giờ", "Đảm bảo chạy 10 phút", "Đảm bảo chạy 1 phút", "Dùng bình xăng 2 lít"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Đường ống hút của máy bơm chữa cháy (Suction Line) bắt buộc phải lắp đặt phụ kiện gì để tránh hiện tượng đọng bọt khí gây sâm thực bơm?",
        "opts": ["Lắp Côn thu lệch tâm (Eccentric Reducer) có mặt phẳng nằm ở phía trên đường ống hút", "Lắp Côn thu đồng tâm", "Lắp van tiết lưu đường kính nhỏ", "Lắp ống uốn cong rủ xuống"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Tủ điều khiển máy bơm chữa cháy (Fire Pump Controller) phải đáp ứng tiêu chuẩn điện bảo vệ và chế độ khởi động bơm như thế nào?",
        "opts": ["Cấp bảo vệ tối thiểu IP54/IP55, có chế độ TỰ ĐỘNG (Auto) và BẰNG TAY (Manual), khởi động sao/tam giác hoặc biến tần/khởi động mềm", "Dùng cầu dao đảo chiều 2 pha", "Chỉ có chế độ điều khiển bằng tay", "Cấp bảo vệ IP10"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Thử tải cụm bơm chữa cháy trong nghiệm thu PCCC công trình (Commissioning Fire Pump Test): Yêu cầu thời gian chạy thử tải liên tục là bao nhiêu?",
        "opts": ["Chạy kiểm tra liên tục 2 giờ (bơm điện và bơm Diesel) đạt 100% và 150% lưu lượng thiết kế, thông số áp suất và nhiệt độ động cơ ổn định", "Chạy thử 2 phút", "Chạy thử 10 giây", "Chỉ cần nhấp nháy động cơ"],
        "correct": 0
    },

    # 5. Hệ thống chữa cháy bằng khí (FM200/CO2/Novec) - 7 câu
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Hệ thống chữa cháy bằng khí sạch (FM200 / Novec 1230 / Nitrogen) thường được ưu tiên thiết kế thi công cho các khu vực nào?",
        "opts": ["Phòng máy chủ Data Center, phòng tổng đài, trung tâm điều khiển, phòng lưu trữ hồ sơ tài liệu quý", "Sân thượng ngoài trời", "Bãi xe ô tô ngoài trời", "Nhà bếp nấu ăn"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Quy trình cảnh báo và trì hoãn thời gian xả khí (Delay Time 30 giây) của hệ thống chữa cháy khí khi có tín hiệu báo cháy kích hoạt nhằm mục đích gì?",
        "opts": ["Cho phép sơ tán toàn bộ người ra khỏi phòng kín và đóng chặt các cửa gió/quạt thông gió trước khi xả khí", "Để khí tự làm mát trong bình", "Để chờ xe chữa cháy đến", "Trì hoãn xả khí 24 giờ"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Thiết bị Công tắc áp lực xả khí (Discharge Pressure Switch) lắp trên đường ống xả khí chữa cháy FM200 có nhiệm vụ gì?",
        "opts": ["Gửi tín hiệu xác nhận khí ĐÃ XẢ về tủ điều khiển trung tâm để bật Đèn cảnh báo 'KHÍ ĐÃ XẢ - CẤM VÀO' ngoài cửa phòng", "Khóa van gas sinh hoạt", "Bật quạt hút khói", "Tắt toàn bộ đèn chiếu sáng"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Nguồn kích hoạt xả khí tự động từ tủ trung tâm chữa cháy khí đến van kích hoạt bình chứa khí (Solenoid Valve) sử dụng điện áp bao nhiêu?",
        "opts": ["Điện áp 24V DC", "Điện áp 380V AC", "Điện áp 110V AC", "Điện áp 220V AC"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Yêu cầu độ kín của phòng (Room Integrity Test) được bảo vệ bằng hệ thống chữa cháy khí FM200/CO2 khi xả khí là gì?",
        "opts": ["Phòng phải kín hoàn toàn, tự động ngắt hệ thống điều hòa thông gió và đóng van ngăn cháy (FD) khi có lệnh xả khí", "Mở toang tất cả các cửa sổ", "Bật quạt hút gió tối đa khi xả khí", "Không cần cửa kín"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Khi chữa cháy bằng khí CO2 trong không gian kín, nguy hiểm lớn nhất đối với kỹ thuật viên và con người là gì?",
        "opts": ["Khí CO2 làm giảm nồng độ Oxy xuống dưới 15% gây ngạt thở cấp tính dẫn đến tử vong nhanh chóng", "Gây dị ứng da nhẹ", "Gây tiếng ồn nhẹ", "Không có nguy hiểm gì"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Nút ấn xả khí khẩn cấp bằng tay (Manual Release Button) và Nút ấn tạm dừng xả khí (Abort Button) phải được bố trí ở đâu?",
        "opts": ["Bố trí bên ngoài lối ra vào cửa chính của phòng được bảo vệ, ở độ cao dễ thao tác 1.3m - 1.4m", "Giấu kín trên trần kỹ thuật", "Đặt bên trong két sắt khóa chặt", "Đặt dưới đáy bình khí"],
        "correct": 0
    },

    # 6. Bình chữa cháy xách tay & Kiểm tra nghiệm thu PCCC - 7 câu
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Bình chữa cháy xách tay dạng bột ABC (như MFZL4, MFZL8) thích hợp dùng để dập tắt các loại đám cháy nào?",
        "opts": ["Đám cháy chất rắn (Class A), chất lỏng (Class B), chất khí (Class C) và thiết bị điện mang điện", "Chỉ dập được cháy kim loại kiềm Na, K", "Chỉ dập cháy nước", "Không dập được cháy điện"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Đồng hồ đo áp suất (Manometer) trên cổ bình chữa cháy bột ABC chỉ vạch màu XANH (Green) báo hiệu trạng thái gì của bình?",
        "opts": ["Áp suất khí nén trong bình đủ tiêu chuẩn làm việc (bình bình thường sẵn sàng sử dụng)", "Bình bị quá áp nguy hiểm", "Bình bị hết áp khí nén (tụt áp)", "Bình bị hỏng vỏ"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Bình chữa cháy xách tay bằng khí CO2 (như MT3, MT5) khi phun dập đám cháy cần lưu ý tuyệt đối điều gì để tránh tai nạn lao động?",
        "opts": ["Không cầm trực tiếp tay vào loa phun hoặc ống nối kim loại để tránh bị bỏng lạnh (-79°C)", "Không được đứng đầu hướng gió", "Phải xịt vào mắt", "Không cần lưu ý"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Khoảng cách di chuyển tối đa từ bất kỳ điểm nào trên mặt bằng công trình đến vị trí đặt bình chữa cháy xách tay nguy hiểm cháy trung bình là bao nhiêu?",
        "opts": ["Không vượt quá 20m (hoặc 15m cho khu vực nguy hiểm cháy cao)", "Không vượt quá 100m", "Không vượt quá 200m", "Chỉ đặt 1 bình ở tầng trệt"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Nghiệm thu tính liên động tự động của hệ thống PCCC khi có tín hiệu báo cháy từ tủ trung tâm (Fire Alarm Interlock Test):",
        "opts": ["Tự động ngắt hệ thống điện hạ thế/điều hòa, tự động hạ cửa cuốn chống cháy, kích hoạt quạt tăng áp giếng thang & quạt hút khói hành lang, gọi thang máy về tầng trệt", "Không tác động gì đến các hệ thống khác", "Chỉ nháy đèn báo", "Tắt toàn bộ nước tòa nhà"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Màu sơn tiêu chuẩn quy định cho đường ống cấp nước chữa cháy, bình chữa cháy và các thiết bị PCCC theo TCVN là màu gì?",
        "opts": ["Màu ĐỎ (Red - RAL 3000)", "Màu XANH LÁ CÂY", "Màu VÀNG KẺ ĐEN", "Màu TRẮNG TÍCH ĐIỆN"],
        "correct": 0
    },
    {
        "cat": "Lý thuyết - Phòng cháy chữa cháy",
        "set": "Lý thuyết - PCCC Bậc 2 & Bậc 3",
        "q": "Hồ sơ nghiệm thu hoàn công PCCC công trình đưa vào hoạt động bắt buộc phải có văn bản pháp lý quan trọng nhất nào do cơ quan Cảnh sát PCCC & CNCH cấp?",
        "opts": ["Văn bản Chấp thuận kết quả nghiệm thu về PCCC của cơ quan Cảnh sát PCCC & CNCH thẩm quyền", "Biên bản họp nội bộ thầu phụ", "Giấy bảo hành của nhà sản xuất bình", "Hóa đơn mua bán vật tư"],
        "correct": 0
    }
]

print(f"Generated {len(pccc_questions)} professional PCCC Theory Questions.")

# 2. UPDATE QUESTIONS.JS ACROSS ALL BUILD PATHS
js_path = r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\questions.js'
with open(js_path, 'r', encoding='utf-8') as f:
    text = f.read()

questions = json.loads(text[text.find('['):text.rfind(']')+1])

# Remove old PCCC theory if any
questions = [q for q in questions if not str(q.get('category', '')).startswith('Lý thuyết - Phòng cháy')]

formatted_pccc = []
for idx, q in enumerate(pccc_questions, start=1):
    q_obj = {
        "id": f"q_pccc_theory_{idx}",
        "type": "multiple_choice",
        "category": q["cat"],
        "exam_set": q["set"],
        "question": q["q"],
        "options": q["opts"],
        "correct_index": q["correct"]
    }
    formatted_pccc.append(q_obj)

questions.extend(formatted_pccc)
print(f"Total questions in bank after adding 50 PCCC theory questions: {len(questions)}.")

new_js = f"// File questions.js - Tích hợp đầy đủ 50 câu hỏi Lý thuyết PCCC chuẩn TCVN & QCVN hiện hành\nconst QUESTIONS = {json.dumps(questions, ensure_ascii=False, indent=2)};\n"

js_paths = [
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\questions.js',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\apk_unpacked\assets\www\questions.js',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\android_build\android\app\src\main\assets\public\questions.js',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\android_build\www\questions.js'
]

for p in js_paths:
    if os.path.exists(os.path.dirname(p)):
        with open(p, 'w', encoding='utf-8') as f:
            f.write(new_js)
        print(f"✅ Saved questions.js to {p}")

# 3. EXPORT EXCEL & WORD BACKUP FILES TO USER'S DESTINATION FOLDER
dest_folder = r'D:\LINH TINH\AI\Khung chương trình đào tạo nghề Điện - CTN\De thi Ý gửi\phần đáp án'
os.makedirs(dest_folder, exist_ok=True)

excel_out = os.path.join(dest_folder, 'Bo_De_Thi_Ly_Thuyet_PCCC_50_Cau_TCVN.xlsx')
word_out = os.path.join(dest_folder, 'Bo_De_Thi_Ly_Thuyet_PCCC_50_Cau_TCVN.docx')

# Create Excel
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Bộ Đề Lý Thuyết PCCC 50 Câu"

headers = ["STT", "Mã Câu Hỏi", "Hạng Mục TCVN", "Nội Dung Câu Hỏi Lý Thuyết PCCC", "Đáp Án A", "Đáp Án B", "Đáp Án C", "Đáp Án D", "Đáp Án Đúng"]
ws.append(headers)

header_fill = PatternFill(start_color="990000", end_color="990000", fill_type="solid")
header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")

for col_num, header in enumerate(headers, 1):
    cell = ws.cell(row=1, column=col_num)
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

for idx, q in enumerate(formatted_pccc, 1):
    row = [
        idx,
        q["id"],
        q["category"],
        q["question"],
        q["options"][0],
        q["options"][1],
        q["options"][2],
        q["options"][3],
        chr(65 + q["correct_index"])
    ]
    ws.append(row)

ws.column_dimensions['A'].width = 6
ws.column_dimensions['B'].width = 18
ws.column_dimensions['C'].width = 30
ws.column_dimensions['D'].width = 65
ws.column_dimensions['E'].width = 32
ws.column_dimensions['F'].width = 32
ws.column_dimensions['G'].width = 32
ws.column_dimensions['H'].width = 32
ws.column_dimensions['I'].width = 12

wb.save(excel_out)
print(f"✅ Exported Excel backup: {excel_out}")

# Create Word
doc = docx.Document()
doc.add_heading('BỘ ĐỀ THI LÝ THUYẾT PHÒNG CHÁY CHỮA CHÁY (50 CÂU)', level=0)
p_sub = doc.add_paragraph('Tiêu chuẩn áp dụng: TCVN 5738:2021 | TCVN 7336:2021 | TCVN 3890:2023 | QCVN 02:2020/BCA | QCVN 06:2022/BXD')
p_sub.runs[0].font.italic = True

for idx, q in enumerate(formatted_pccc, 1):
    doc.add_heading(f"Câu {idx}: {q['question']}", level=2)
    for opt_idx, opt in enumerate(q["options"]):
        prefix = f"  {chr(65 + opt_idx)}. "
        p_opt = doc.add_paragraph(prefix + opt)
        if opt_idx == q["correct_index"]:
            p_opt.runs[0].font.bold = True
            p_opt.runs[0].font.color.rgb = RGBColor(180, 0, 0)
    doc.add_paragraph(f"👉 ĐÁP ÁN ĐÚNG: {chr(65 + q['correct_index'])}")
    doc.add_paragraph("-" * 50)

doc.save(word_out)
print(f"✅ Exported Word backup: {word_out}")
