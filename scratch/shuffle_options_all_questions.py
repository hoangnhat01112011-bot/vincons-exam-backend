import json
import random
import os
import shutil
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
import docx
from docx.shared import Pt, RGBColor

# Set seed for reproducible pseudo-random distribution
random.seed(20260826)

js_path = r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\questions.js'
with open(js_path, 'r', encoding='utf-8') as f:
    text = f.read()

questions = json.loads(text[text.find('['):text.rfind(']')+1])

print(f"Total questions before shuffling: {len(questions)}")

# Distribution counter
distrib = {0: 0, 1: 0, 2: 0, 3: 0}

for idx, q in enumerate(questions):
    opts = q.get("options", [])
    corr_idx = q.get("correct_index", 0)

    if opts and 0 <= corr_idx < len(opts):
        correct_text = opts[corr_idx]
        
        # Shuffle options array
        random.shuffle(opts)
        
        # New correct index
        new_corr_idx = opts.index(correct_text)
        q["options"] = opts
        q["correct_index"] = new_corr_idx
        distrib[new_corr_idx] += 1

print(f"✅ Shuffled correct answer distribution across bank:")
print(f"   Option A (Index 0): {distrib[0]} questions ({distrib[0]/len(questions)*100:.1f}%)")
print(f"   Option B (Index 1): {distrib[1]} questions ({distrib[1]/len(questions)*100:.1f}%)")
print(f"   Option C (Index 2): {distrib[2]} questions ({distrib[2]/len(questions)*100:.1f}%)")
print(f"   Option D (Index 3): {distrib[3]} questions ({distrib[3]/len(questions)*100:.1f}%)")

# Save updated questions.js
new_js = f"// File questions.js - Đã xáo trộn ngẫu nhiên đáp án A, B, C, D phân bổ đều giữa các câu hỏi\nconst QUESTIONS = {json.dumps(questions, ensure_ascii=False, indent=2)};\n"

js_paths = [
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\questions.js',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\apk_unpacked\assets\www\questions.js',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\android_build\android\app\src\main\assets\public\questions.js',
    r'D:\LINH TINH\AI\PHAN MEM\vincons-test-app\android_build\www\questions.js'
]

for p in js_paths:
    if os.path.exists(os.path.dirname(p)):
        with open(p, 'w', encoding='utf-8') as f:
            f.write(new_js)
        print(f"✅ Saved questions.js to {p}")

# Re-export PCCC Word and Excel backups with updated shuffled correct answers
dest_folder = r'D:\LINH TINH\AI\Khung chương trình đào tạo nghề Điện - CTN\De thi Ý gửi\phần đáp án'
os.makedirs(dest_folder, exist_ok=True)

excel_out = os.path.join(dest_folder, 'Bo_De_Thi_Ly_Thuyet_PCCC_50_Cau_TCVN.xlsx')
word_out = os.path.join(dest_folder, 'Bo_De_Thi_Ly_Thuyet_PCCC_50_Cau_TCVN.docx')

pccc_qs = [q for q in questions if str(q.get('category', '')).startswith('Lý thuyết - Phòng cháy')]

if pccc_qs:
    # Update Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Bộ Đề Lý Thuyết PCCC 50 Câu"
    headers = ["STT", "Mã Câu Hỏi", "Hạng Mục TCVN", "Nội Dung Câu Hỏi Lý Thuyết PCCC", "Đáp Án A", "Đáp Án B", "Đáp Án C", "Đáp Án D", "Đáp Án Đúng"]
    ws.append(headers)

    header_fill = PatternFill(start_color="990000", end_color="990000", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for idx, q in enumerate(pccc_qs, 1):
        row = [
            idx,
            q["id"],
            q["category"],
            q["question"],
            q["options"][0],
            q["options"][1],
            q["options"][2],
            q["options"][3],
            chr(65 + q["correct_index"])
        ]
        ws.append(row)

    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 30
    ws.column_dimensions['D'].width = 65
    ws.column_dimensions['E'].width = 32
    ws.column_dimensions['F'].width = 32
    ws.column_dimensions['G'].width = 32
    ws.column_dimensions['H'].width = 32
    ws.column_dimensions['I'].width = 12
    wb.save(excel_out)
    print(f"✅ Re-exported Excel backup with shuffled answers: {excel_out}")

    # Update Word
    doc = docx.Document()
    doc.add_heading('BỘ ĐỀ THI LÝ THUYẾT PHÒNG CHÁY CHỮA CHÁY (50 CÂU)', level=0)
    p_sub = doc.add_paragraph('Tiêu chuẩn áp dụng: TCVN 5738:2021 | TCVN 7336:2021 | TCVN 3890:2023 | QCVN 02:2020/BCA | QCVN 06:2022/BXD')
    p_sub.runs[0].font.italic = True

    for idx, q in enumerate(pccc_qs, 1):
        doc.add_heading(f"Câu {idx}: {q['question']}", level=2)
        for opt_idx, opt in enumerate(q["options"]):
            prefix = f"  {chr(65 + opt_idx)}. "
            p_opt = doc.add_paragraph(prefix + opt)
            if opt_idx == q["correct_index"]:
                p_opt.runs[0].font.bold = True
                p_opt.runs[0].font.color.rgb = RGBColor(180, 0, 0)
        doc.add_paragraph(f"👉 ĐÁP ÁN ĐÚNG: {chr(65 + q['correct_index'])}")
        doc.add_paragraph("-" * 50)

    doc.save(word_out)
    print(f"✅ Re-exported Word backup with shuffled answers: {word_out}")
